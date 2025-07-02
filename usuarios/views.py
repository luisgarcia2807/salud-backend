from rest_framework import viewsets, serializers, status
from .models import Alergia, Doctor, DoctorCentro, EnfermedadComun, EnfermedadPersistente, ExamenLabImagenologia, ExamenLaboratorio, GrupoSanguineo, MedicamentoCronico, Paciente, PacienteAlergia, PacienteEnfermedadPersistente, PacienteMedicamentoCronico, PerfilBebe, RegistroVacuna, Usuario, Vacuna
from .serializers import AlergiaSerializer, DoctorPacienteDetalleSerializer, DoctorSerializer, DocumentoEscaneadoSerializer, EnfermedadComunSerializer, EnfermedadPersistenteSerializer, ExamenImagenologiaSerializer, ExamenLaboratorioSerializer, GrupoSanguineoSerializer, HistoriaClinicaPacienteSerializer, MedicamentoCronicoSerializer, PacienteAlergiaSerializer, PacienteEnfermedadComunSerializer, PacienteEnfermedadPersistenteSerializer, PacienteMedicamentoCronicoSerializer, PacienteSerializer, PerfilBebeRegistroSerializer, RegistroVacunaSerializer, UsuarioSerializer, VacunaSerializer
from rest_framework.permissions import AllowAny
from rest_framework_simplejwt.views import TokenObtainPairView, TokenRefreshView
from django.contrib.auth import authenticate
from rest_framework.response import Response
from rest_framework.views import APIView
from .models import CentroMedico
from .serializers import CentroMedicoSerializer
from .models import Especialidad
from .serializers import EspecialidadSerializer
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework import status
from .models import PruebaImagen
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from .models import DocumentoEscaneado
from .serializers import DocumentoEscaneadoSerializer



class UsuarioDetailView(APIView):
    def get(self, request, id_usuario, *args, **kwargs):
        try:
            # Buscar usuario usando 'id_usuario'
            usuario = Usuario.objects.get(id_usuario=id_usuario)
            serializer = UsuarioSerializer(usuario)
            return Response(serializer.data, status=status.HTTP_200_OK)
        except Usuario.DoesNotExist:
            return Response({"detail": "Usuario no encontrado."}, status=status.HTTP_404_NOT_FOUND)


class PerfilBebeDetailView(APIView):
    def get(self, request, id_bebe, *args, **kwargs):
        try:
            bebe = PerfilBebe.objects.get(id=id_bebe)
            data = {
                "nombre": bebe.nombre,
                "apellido": bebe.apellido,
                "fecha_nacimiento": bebe.fecha_nacimiento,
                "sexo": bebe.sexo,
            }
            return Response(data, status=status.HTTP_200_OK)
        except PerfilBebe.DoesNotExist:
            return Response({"detail": "Perfil bebé no encontrado."}, status=status.HTTP_404_NOT_FOUND)

class PerfilBebeViewSet(viewsets.ModelViewSet):
    queryset = PerfilBebe.objects.all()
    serializer_class = PerfilBebeRegistroSerializer

from rest_framework.generics import ListAPIView

class BebesPorResponsableView(ListAPIView):
    serializer_class = PerfilBebeRegistroSerializer

    def get_queryset(self):
        responsable_id = self.kwargs.get('responsable_id')
        return PerfilBebe.objects.filter(responsable_id=responsable_id)

# ViewSet para el manejo de usuarios
class UsuarioViewSet(viewsets.ModelViewSet):
    queryset = Usuario.objects.all()  # Obtiene todos los usuarios
    serializer_class = UsuarioSerializer  # Usa el serializador de usuarios
    permission_classes = [AllowAny]  # Permite que cualquiera acceda a este ViewSet (solo para pruebas)
    
    def create(self, request, *args, **kwargs):
        password = request.data.get('password')  # Obtener la contraseña en texto plano

        # Si necesitas lógica adicional antes de crear, como encriptar la contraseña
        return super().create(request, *args, **kwargs)

# Vista personalizada para obtener el token JWT
class CustomTokenObtainPairView(TokenObtainPairView):
    permission_classes = [AllowAny]  # Permitir que cualquier usuario acceda para pruebas

    def post(self, request, *args, **kwargs):
        # Obtiene el email y la contraseña del request
        email = request.data.get('email')
        password = request.data.get('password')

        # Verifica si se proporcionaron el email y la contraseña
        if not email or not password:
            return Response({"detail": "Debe proporcionar un email y una contraseña."},
                             status=status.HTTP_400_BAD_REQUEST)

        # Autenticar al usuario usando el email (en vez del username)
        user = authenticate(username=email, password=password)

        # Si las credenciales no son válidas
        if user is None:
            # Verificamos si el email existe
            try:
                user = Usuario.objects.get(email=email)
            except Usuario.DoesNotExist:
                # Si el usuario no existe, el error es por el email
                return Response({"detail": "El email no está registrado."},
                                 status=status.HTTP_404_NOT_FOUND)

            # Si el email existe pero la contraseña es incorrecta
            return Response({"detail": "Contraseña incorrecta. Intenta de nuevo."},
                             status=status.HTTP_401_UNAUTHORIZED)

        # Si las credenciales son correctas, generar el token
        from rest_framework_simplejwt.tokens import RefreshToken

        refresh = RefreshToken.for_user(user)
        access_token = str(refresh.access_token)

        # Devolver el token y la información adicional del usuario
        return Response({
            'access': access_token,
            'refresh': str(refresh),
            'id_usuario': user.id_usuario,
            'email': user.email,
            'nombre': user.nombre,
            'apellido': user.apellido
        }, status=status.HTTP_200_OK)

# Vista para refrescar el token JWT
class CustomTokenRefreshView(TokenRefreshView):
    permission_classes = [AllowAny]  # Permitir que cualquier usuario acceda para pruebas

class CentroMedicoListView(APIView):
    def get(self, request):
        centros_medicos = CentroMedico.objects.all()  # Obtenemos todos los centros médicos
        serializer = CentroMedicoSerializer(centros_medicos, many=True)
        return Response(serializer.data)
    

class EspecialidadListView(APIView):
    def get(self, request):
        especialidades = Especialidad.objects.all()  # Obtenemos todas las especialidades
        serializer = EspecialidadSerializer(especialidades, many=True)
        return Response(serializer.data)

class GrupoSanguineoListView(APIView):
    def get(self, request):
        
        grupos_sanguineos = GrupoSanguineo.objects.all()
        serializer = GrupoSanguineoSerializer(grupos_sanguineos, many=True)
        return Response(serializer.data)


class PacientePorUsuarioView(APIView):
    def get(self, request, id_usuario):
        try:
            paciente = Paciente.objects.select_related('id_usuario', 'id_sangre').get(id_usuario=id_usuario)
            serializer = PacienteSerializer(paciente)
            return Response(serializer.data, status=status.HTTP_200_OK)
        except Paciente.DoesNotExist:
            return Response({"detail": "No se encontró paciente asociado a ese usuario."}, status=status.HTTP_404_NOT_FOUND)
        
class PacientePorPerfilBebeView(APIView):
    def get(self, request, id_perfil_bebe):
        try:
            paciente = Paciente.objects.select_related('perfil_bebe', 'id_sangre').get(perfil_bebe=id_perfil_bebe)
            serializer = PacienteSerializer(paciente)
            return Response(serializer.data, status=status.HTTP_200_OK)
        except Paciente.DoesNotExist:
            return Response({"detail": "No se encontró paciente asociado a ese perfil de bebé."}, status=status.HTTP_404_NOT_FOUND)


class DatosBasicosPorPacienteView(APIView):
    def get(self, request, id_paciente):
        try:
            paciente = Paciente.objects.select_related('id_usuario', 'perfil_bebe', 'id_sangre').get(id_paciente=id_paciente)

            data = {
                "id_paciente": paciente.id_paciente,
                "id_sangre": paciente.id_sangre.id_sangre if paciente.id_sangre else None,
                "tipo_sangre": paciente.id_sangre.tipo_sangre if paciente.id_sangre else None,
                "token": paciente.token,
            }

            if paciente.id_usuario:
                data["tipo"] = "usuario"
                data["id_u"] = paciente.id_usuario.id_usuario
            elif paciente.perfil_bebe:
                data["tipo"] = "bebe"
                data["id_u"] = paciente.perfil_bebe.id
            else:
                return Response(
                    {"detail": "El paciente no tiene usuario ni perfil de bebé asociado."},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            return Response(data, status=status.HTTP_200_OK)

        except Paciente.DoesNotExist:
            return Response({"detail": "Paciente no encontrado."}, status=status.HTTP_404_NOT_FOUND)


class DatosBasicosPorTokenView(APIView):
    def get(self, request, token):
        try:
            paciente = Paciente.objects.select_related('id_usuario', 'perfil_bebe', 'id_sangre').get(token=token)

            data = {
                "id_paciente": paciente.id_paciente,
                "id_sangre": paciente.id_sangre.id_sangre if paciente.id_sangre else None,
                "tipo_sangre": paciente.id_sangre.tipo_sangre if paciente.id_sangre else None,
                "token": paciente.token,
            }

            if paciente.id_usuario:
                data["tipo"] = "usuario"
                data["id_u"] = paciente.id_usuario.id_usuario
                data["nombre"] = paciente.id_usuario.nombre
                data["apellido"] = paciente.id_usuario.apellido

            elif paciente.perfil_bebe:
                data["tipo"] = "bebe"
                data["id_u"] = paciente.perfil_bebe.id
                data["nombre"] = paciente.perfil_bebe.nombre
                data["apellido"] = paciente.perfil_bebe.apellido

            else:
                return Response(
                    {"detail": "El paciente no tiene usuario ni perfil de bebé asociado."},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            return Response(data, status=status.HTTP_200_OK)

        except Paciente.DoesNotExist:
            return Response({"detail": "Paciente no encontrado con ese token."}, status=status.HTTP_404_NOT_FOUND)

class ListaAlergias(APIView):
    def get(self, request):
        tipo = request.GET.get('tipo', None)  # Obtiene el parámetro 'tipo' de la URL

        if tipo:
            # Filtra las alergias por el tipo especificado
            alergias = Alergia.objects.filter(tipo=tipo)
        else:
            # Si no se proporciona el parámetro 'tipo', retorna todas las alergias
            alergias = Alergia.objects.all()

        serializer = AlergiaSerializer(alergias, many=True)
        return Response(serializer.data)
    
class PacienteAlergiaViewSet(viewsets.ModelViewSet):
    queryset = PacienteAlergia.objects.all()
    serializer_class = PacienteAlergiaSerializer

class AlergiasPorPacienteView(APIView):
    def get(self, request, id_paciente):
        tipo = request.query_params.get('tipo')  # Ejemplo: tipo=alimentaria

        alergias = PacienteAlergia.objects.filter(paciente_id=id_paciente).select_related('alergia', 'doctor_aprobador__id_usuario')

        if tipo:
            alergias = alergias.filter(alergia__tipo=tipo)  # el tipo debe coincidir con el valor interno del Enum

        resultados = []
        for a in alergias:
            resultados.append({
                'id': a.id,
                'nombre_alergia': a.alergia.nombre,
                'tipo_alergia': a.alergia.get_tipo_display(),
                'gravedad': a.get_gravedad_display(),
                'observacion': a.observacion,
                'aprobado': a.aprobado,
                'doctor_aprobador': (
                    f"{a.doctor_aprobador.id_usuario.nombre} {a.doctor_aprobador.id_usuario.apellido}"
                    if a.doctor_aprobador else None
                )
            })

        return Response(resultados, status=status.HTTP_200_OK)

class VacunaListView(APIView):
    def get(self, request):
        vacunas = Vacuna.objects.all()  # Obtener todas las vacunas
        serializer = VacunaSerializer(vacunas, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)


class VacunaPacienteViewSet(viewsets.ModelViewSet):
    
   
    queryset = RegistroVacuna.objects.all()
    serializer_class =RegistroVacunaSerializer


class VacunasPorPacienteView(APIView):
    def get(self, request, id_paciente):
        vacunas_paciente = RegistroVacuna.objects.filter(paciente_id=id_paciente).select_related('vacuna','doctor_aprobador__id_usuario')
        
        resultados = []
        for vp in vacunas_paciente:
            resultados.append({
                'id': vp.id, 
                'nombre_vacuna': vp.vacuna.nombre,
                'descripcion_vacuna': vp.vacuna.descripcion,
                'max_dosis': vp.vacuna.max_dosis,
                'fecha_aplicacion': vp.fecha_aplicacion,
                'dosis': vp.dosis,
                'observacion': vp.observacion,
                'aprobado': vp.aprobado,
                'doctor_aprobador': (
                    f"{vp.doctor_aprobador.id_usuario.nombre} {vp.doctor_aprobador.id_usuario.apellido}"
                    if vp.doctor_aprobador else None
                )
            })

        return Response(resultados, status=status.HTTP_200_OK)
    
class EnfermedadPersistenteListView(APIView):
    def get(self, request):
        tipo = request.GET.get('tipo', None)  # Obtener el parámetro 'tipo' de la URL

        if tipo:
            # Filtra las enfermedades persistentes por tipo
            enfermedades = EnfermedadPersistente.objects.filter(tipo=tipo)
        else:
            # Si no se proporciona 'tipo', retorna todas las enfermedades
            enfermedades = EnfermedadPersistente.objects.all()

        serializer = EnfermedadPersistenteSerializer(enfermedades, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)


class PacienteEnfermedadPersistenteViewSet(viewsets.ModelViewSet):
    queryset = PacienteEnfermedadPersistente.objects.all()  # Obtiene todos los registros de enfermedades persistentes
    serializer_class = PacienteEnfermedadPersistenteSerializer  # Usa el serializador de PacienteEnfermedadPersistente


class EnfermedadesPorPacienteView(APIView):
    def get(self, request, id_paciente):
        tipo_param = request.query_params.get('tipo')  # Captura el tipo desde la URL

        enfermedades = PacienteEnfermedadPersistente.objects.filter(
            paciente_id=id_paciente
        ).select_related('enfermedad', 'doctor_aprobador__id_usuario')

        if tipo_param:
            enfermedades = enfermedades.filter(enfermedad__tipo=tipo_param)

        resultados = []
        for e in enfermedades:
            resultados.append({
                'id': e.id,
                'nombre_enfermedad': e.enfermedad.nombre,
                'Tipo_enfermedad': e.enfermedad.get_tipo_display(),
                'fecha_diagnostico': e.fecha_diagnostico,
                'observacion': e.observacion,
                'aprobado': e.aprobado,
                'doctor_aprobador': (
                    f"{e.doctor_aprobador.id_usuario.nombre} {e.doctor_aprobador.id_usuario.apellido}"
                    if e.doctor_aprobador else None
                )
            })

        return Response(resultados, status=status.HTTP_200_OK)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.db.models import Q
from .models import PacienteEnfermedadComun

class EnfermedadComunListView(APIView):
    def get(self, request):
        tipo = request.query_params.get('tipo', None)
        enfermedades = EnfermedadComun.objects.all()
        if tipo:
            enfermedades = enfermedades.filter(tipo=tipo)
        serializer = EnfermedadComunSerializer(enfermedades, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
class PacienteEnfermedadComunViewSet(viewsets.ModelViewSet):
    queryset = PacienteEnfermedadComun.objects.all()
    serializer_class = PacienteEnfermedadComunSerializer

class EnfermedadesComunesPorPacienteView(APIView):
    def get(self, request, id_paciente):
        activas = request.query_params.get('activas')
        tipo = request.query_params.get('tipo')
        nombre = request.query_params.get('nombre')

        enfermedades = PacienteEnfermedadComun.objects.filter(paciente_id=id_paciente).select_related('enfermedad', 'doctor_aprobador__id_usuario')

        if activas:
            if activas.lower() == 'true':
                enfermedades = enfermedades.filter(fecha_recuperacion__isnull=True)
            elif activas.lower() == 'false':
                enfermedades = enfermedades.filter(fecha_recuperacion__isnull=False)

        if tipo:
            enfermedades = enfermedades.filter(enfermedad__tipo=tipo)

        if nombre:
            enfermedades = enfermedades.filter(enfermedad__nombre__icontains=nombre)

        resultados = []
        for e in enfermedades:
            resultados.append({
                'id': e.id,
                'nombre_enfermedad': e.enfermedad.nombre,
                'tipo_enfermedad': e.enfermedad.get_tipo_display(),
                'descripcion': e.enfermedad.descripcion,
                'fecha_diagnostico': e.fecha_diagnostico,
                'fecha_recuperacion': e.fecha_recuperacion,
                'observacion': e.observacion,
                'aprobado': e.aprobado,
                'doctor_aprobador': (
                    f"{e.doctor_aprobador.id_usuario.nombre} {e.doctor_aprobador.id_usuario.apellido}"
                    if e.doctor_aprobador else None
                )
            })

        return Response(resultados, status=status.HTTP_200_OK)


class DoctoresActivosInactivosView(APIView):
    def get(self, request, id_centromedico):
        # Obtener todos los doctores asociados a este centro médico
        doctores = Doctor.objects.filter(
            id_doctor__in=DoctorCentro.objects.filter(id_centromedico=id_centromedico).values('id_doctor')
        )
        
        # Serializamos la lista de doctores
        serializer = DoctorSerializer(doctores, many=True)
        
        return Response(serializer.data, status=status.HTTP_200_OK)
    
class ActivarDoctorView(APIView):
    def post(self, request, id_centromedico, id_doctor):
        # Buscar si ya existe una solicitud de DoctorCentro
        try:
            doctor_centro = DoctorCentro.objects.get(id_doctor=id_doctor, id_centromedico=id_centromedico)
        except DoctorCentro.DoesNotExist:
            return Response({"detail": "No se encontró la solicitud del doctor."}, status=status.HTTP_404_NOT_FOUND)

        # Si la solicitud está pendiente, activarla
        if not doctor_centro.aceptado_por_centromedico:
            doctor_centro.aceptado_por_centromedico = True
            doctor_centro.save()
            return Response({"detail": "Doctor activado correctamente."}, status=status.HTTP_200_OK)
        
        return Response({"detail": "El doctor ya está activo en este centro médico."}, status=status.HTTP_400_BAD_REQUEST)
    
class SubirImagenPruebaView(APIView):
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        nombre = request.data.get('nombre')
        imagen = request.FILES.get('imagen')

        print("Datos recibidos: ", request.data)  # Para verificar los datos
        print("Archivos recibidos: ", request.FILES)  # Para verificar los archivos

        if not nombre or not imagen:
            return Response({'error': 'Nombre e imagen son requeridos'}, status=status.HTTP_400_BAD_REQUEST)

        nueva_imagen = PruebaImagen(nombre=nombre, imagen=imagen)
        nueva_imagen.save()

        return Response({'mensaje': 'Imagen subida correctamente'}, status=status.HTTP_201_CREATED)

    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        nombre = request.data.get('nombre')
        imagen = request.FILES.get('imagen')

        if not nombre or not imagen:
            return Response({'error': 'Nombre e imagen son requeridos'}, status=status.HTTP_400_BAD_REQUEST)

        nueva_imagen = PruebaImagen(nombre=nombre, imagen=imagen)
        nueva_imagen.save()

        return Response({'mensaje': 'Imagen subida correctamente'}, status=status.HTTP_201_CREATED)
    
@api_view(['PUT'])
def actualizar_foto_perfil(request, id_usuario):
    try:
        usuario = Usuario.objects.get(id_usuario=id_usuario)
    except Usuario.DoesNotExist:
        return Response({'error': 'Usuario no encontrado.'}, status=status.HTTP_404_NOT_FOUND)

    foto = request.FILES.get('foto_perfil')
    if not foto:
        return Response({'error': 'Debe enviar una imagen con el campo "foto_perfil".'}, status=status.HTTP_400_BAD_REQUEST)

    usuario.foto_perfil = foto
    usuario.save()
    return Response({'mensaje': 'Foto de perfil actualizada correctamente.'}, status=status.HTTP_200_OK)



class MedicamentoCronicoListView(APIView):
    def get(self, request):
        # Obtenemos todos los medicamentos crónicos
        medicamentos = MedicamentoCronico.objects.all()
        # Serializamos los medicamentos
        serializer = MedicamentoCronicoSerializer(medicamentos, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)

class PacienteMedicamentoCronicoViewSet(viewsets.ModelViewSet):
    queryset = PacienteMedicamentoCronico.objects.all()
    serializer_class = PacienteMedicamentoCronicoSerializer

class TratamientosCronicosPorPacienteView(APIView):
    def get(self, request, id_paciente):
        tratamientos = PacienteMedicamentoCronico.objects.filter(
            id_paciente=id_paciente
        ).select_related('id_medicamento_cronico', 'doctor_aprobador__id_usuario')

        resultados = []
        for t in tratamientos:
            resultados.append({
                'id':t.id,
                'nombre_medicamento': t.id_medicamento_cronico.nombre,
                'descripcion_medicamento': t.id_medicamento_cronico.descripcion,
                'fecha_inicio': t.fecha_inicio,
                'dosis': t.dosis,
                'frecuencia': t.frecuencia,
                'observaciones': t.observaciones,
                'aprobado': t.aprobado,
                'doctor_aprobador': (
                    f"{t.doctor_aprobador.id_usuario.nombre} {t.doctor_aprobador.id_usuario.apellido}"
                    if t.doctor_aprobador else None
                )
            })

        return Response(resultados, status=status.HTTP_200_OK)

class SubirDocumentoView(APIView):
    parser_classes = (MultiPartParser, FormParser)
    def post(self, request, *args, **kwargs):
        serializer = DocumentoEscaneadoSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"mensaje": "PDF subido correctamente", "data": serializer.data})
        return Response(serializer.errors, status=400)
    


from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework import status
from .serializers import ExamenLaboratorioSerializer
from .models import ExamenLaboratorio
from .supabase_config import supabase


class ExamenLaboratorioView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def get(self, request, paciente_id):
        examenes = ExamenLaboratorio.objects.filter(paciente_id=paciente_id).select_related('doctor__id_usuario').order_by('-fecha_subida')
        resultados = []

        for examen in examenes:
            resultados.append({
                'id': examen.id,
                'tipo': examen.tipo,
                'categoria': examen.categoria,
                'nombre_examen': examen.nombre_examen,
                'descripcion': examen.descripcion,
                'fecha_realizacion': examen.fecha_realizacion,
                'archivo': examen.archivo,
                'fecha_subida': examen.fecha_subida,
                'paciente': examen.paciente_id,
                'doctor': examen.doctor_id,
                'nombre_doctor': (
                    f"{examen.doctor.id_usuario.nombre} {examen.doctor.id_usuario.apellido}"
                    if examen.doctor and examen.doctor.id_usuario else None
                )
            })

        return Response(resultados, status=status.HTTP_200_OK)

    def post(self, request):
        archivo = request.FILES.get('archivo')
        if archivo:
            contenido = archivo.read()
            ruta = f"examenes/{archivo.name}"
            supabase.storage.from_("examenes").upload(ruta, contenido, {"content-type": archivo.content_type})
            url_publica = supabase.storage.from_("examenes").get_public_url(ruta)

            data = request.data.dict()
            data["archivo"] = url_publica

            serializer = ExamenLaboratorioSerializer(data=data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        return Response({"error": "Archivo no recibido"}, status=400)
    
    def delete(self, request, examen_id):
        try:
            examen = ExamenLaboratorio.objects.get(id=examen_id)
            examen.delete()
            return Response({"mensaje": "Examen eliminado correctamente"}, status=status.HTTP_204_NO_CONTENT)
        except ExamenLaboratorio.DoesNotExist:
            return Response({"error": "Examen no encontrado"}, status=status.HTTP_404_NOT_FOUND)
    

class ExamenlabImagenologiaView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def get(self, request, paciente_id):
        examenes = ExamenLabImagenologia.objects.filter(paciente_id=paciente_id).select_related('doctor__id_usuario').order_by('-fecha_subida')
        resultados = []

        for examen in examenes:
            resultados.append({
                'id': examen.id,
                'tipo': examen.tipo,
                'categoria': examen.categoria,
                'nombre_examen': examen.nombre_examen,
                'descripcion': examen.descripcion,
                'fecha_realizacion': examen.fecha_realizacion,
                'archivo': examen.archivo,
                'fecha_subida': examen.fecha_subida,
                'paciente': examen.paciente_id,
                'doctor': examen.doctor_id,
                'nombre_doctor': (
                    f"{examen.doctor.id_usuario.nombre} {examen.doctor.id_usuario.apellido}"
                    if examen.doctor and examen.doctor.id_usuario else None
                )
            })

        return Response(resultados, status=status.HTTP_200_OK)

    def post(self, request):
        archivo = request.FILES.get('archivo')
        if archivo:
            contenido = archivo.read()
            ruta = f"imagenologia/{archivo.name}"
            supabase.storage.from_("examenes").upload(ruta, contenido, {"content-type": archivo.content_type})
            url_publica = supabase.storage.from_("examenes").get_public_url(ruta)

            data = request.data.dict()
            data["archivo"] = url_publica

            serializer = ExamenImagenologiaSerializer(data=data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        return Response({"error": "Archivo no recibido"}, status=400)

    def delete(self, request, examen_id):
        try:
            examen = ExamenLabImagenologia.objects.get(id=examen_id)
            examen.delete()
            
            return Response({"mensaje": "Examen eliminado correctamente"}, status=status.HTTP_204_NO_CONTENT)
        except ExamenLabImagenologia.DoesNotExist:
            return Response({"error": "Examen no encontrado"}, status=status.HTTP_404_NOT_FOUND)


import cv2
import numpy as np
from django.http import HttpResponse
from django.views.decorators.csrf import csrf_exempt

@csrf_exempt
def procesar_documento(request):
    if request.method != 'POST':
        return HttpResponse("Método no permitido", status=405)

    if 'imagen' not in request.FILES:
        return HttpResponse("No se envió imagen", status=400)

    file = request.FILES['imagen']
    image_bytes = file.read()

    np_arr = np.frombuffer(image_bytes, np.uint8)
    image = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)

    if image is None:
        return HttpResponse("Imagen inválida o corrupta", status=400)

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    # Aplicar CLAHE para mejorar contraste local y evitar pérdida de detalles
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
    enhanced = clahe.apply(gray)

    # Umbral binario con un valor bajo para atrapar detalles finos
    _, thresh = cv2.threshold(enhanced, 90, 255, cv2.THRESH_BINARY)

    # Invertir para que letras sean negras y fondo blanco
    thresh = cv2.bitwise_not(thresh)

    # Aplicar filtro mediana para reducir ruido manteniendo bordes
    denoised = cv2.medianBlur(thresh, 3)

    # Dilatar suavemente para reforzar letras sin perder detalle
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2,2))
    dilated = cv2.dilate(denoised, kernel, iterations=1)

    # Invertir nuevamente para que letras negras y fondo blanco
    result = cv2.bitwise_not(dilated)

    _, jpeg = cv2.imencode('.jpg', result)
    return HttpResponse(jpeg.tobytes(), content_type="image/jpeg")


from rest_framework.decorators import api_view


@api_view(['GET'])
def proxima_dosis(request, paciente_id, vacuna_id):
    try:
        vacuna = Vacuna.objects.get(id=vacuna_id)
    except Vacuna.DoesNotExist:
        return Response({"detail": "Vacuna no encontrada."}, status=status.HTTP_404_NOT_FOUND)

    registros = RegistroVacuna.objects.filter(paciente_id=paciente_id, vacuna_id=vacuna_id).order_by('dosis')

    if registros.exists():
        ultima_dosis = registros.last().dosis
        if ultima_dosis >= vacuna.max_dosis:
            return Response({"detail": "Ya se han aplicado todas las dosis."}, status=status.HTTP_400_BAD_REQUEST)
        return Response({"proxima_dosis": ultima_dosis + 1}, status=status.HTTP_200_OK)
    
    return Response({"proxima_dosis": 1}, status=status.HTTP_200_OK)



from django.db.models import Max

@api_view(['GET'])
def ultimas_dosis_por_paciente(request, paciente_id):
    # Obtener la última dosis por cada vacuna
    subquery = RegistroVacuna.objects.filter(paciente_id=paciente_id)
    ultimas = subquery.values('vacuna').annotate(ultima_dosis=Max('dosis'))

    resultado = []
    for entry in ultimas:
        vacuna_id = entry['vacuna']
        dosis = entry['ultima_dosis']
        registro = RegistroVacuna.objects.select_related('vacuna', 'doctor_aprobador__id_usuario').get(
            paciente_id=paciente_id,
            vacuna_id=vacuna_id,
            dosis=dosis
        )

        resultado.append({
            'id': registro.id,
            'nombre_vacuna': registro.vacuna.nombre,
            'descripcion_vacuna': registro.vacuna.descripcion,
            'max_dosis': registro.vacuna.max_dosis,
            'fecha_aplicacion': registro.fecha_aplicacion,
            'dosis': registro.dosis,
            'observacion': registro.observacion,
            'aprobado': registro.aprobado,
            'doctor_aprobador': (
                f"{registro.doctor_aprobador.id_usuario.nombre} {registro.doctor_aprobador.id_usuario.apellido}"
                if registro.doctor_aprobador else None
            )
        })

    return Response(resultado)


from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from .models import TratamientoActual, SeguimientoTratamiento
from .serializers import TratamientoActualSerializer, SeguimientoTratamientoSerializer

@api_view(['GET'])
def tratamientos_por_paciente(request, paciente_id):
    tratamientos = TratamientoActual.objects.filter(paciente_id=paciente_id).select_related('medicamento').order_by('-created_at')

    resultados = []
    for t in tratamientos:
        resultados.append({
            'id': t.id,
            'descripcion': t.descripcion,
            'fecha_inicio': t.fecha_inicio,
            'fecha_fin': t.fecha_fin,
            'finalizado': t.finalizado,
            'frecuencia': t.frecuencia,
            'created_at': t.created_at,
            'updated_at': t.updated_at,
            'paciente': t.paciente_id,
            'medicamento': t.medicamento_id,
            'nombre_medicamento': t.medicamento.nombre_comercial if t.medicamento else None,
            'dosis':t.medicamento.concentracion if t.medicamento else None,
            'via': t.medicamento.via_administracion if t.medicamento else None,
            'doctor': t.doctor_id,
            'nombre_doctor': (
                f"{t.doctor.id_usuario.nombre} {t.doctor.id_usuario.apellido}"
                if t.doctor and t.doctor.id_usuario else None
            )
        })

    return Response(resultados)



@api_view(['POST'])
def crear_tratamiento(request):
    paciente_id = request.data.get('paciente')
    medicamento_id = request.data.get('medicamento')

    # Validación mínima necesaria para la restricción
    if not paciente_id or not medicamento_id:
        return Response({'error': 'El paciente y el medicamento son obligatorios para validar duplicación.'},
                        status=status.HTTP_400_BAD_REQUEST)

    # Restricción: no permitir duplicados no finalizados
    tratamiento_existente = TratamientoActual.objects.filter(
        paciente_id=paciente_id,
        medicamento_id=medicamento_id,
        finalizado=False
    ).exists()

    if tratamiento_existente:
        return Response(
            {'error': 'Ya existe un tratamiento activo con este medicamento para el paciente.'},
            status=status.HTTP_400_BAD_REQUEST
        )

    # Continuar con la creación
    serializer = TratamientoActualSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['PATCH'])
def actualizar_tratamiento(request, pk):
    try:
        tratamiento = TratamientoActual.objects.get(pk=pk)
    except TratamientoActual.DoesNotExist:
        return Response({'error': 'Tratamiento no encontrado.'}, status=status.HTTP_404_NOT_FOUND)

    # Solo permitimos actualizar doctor, fecha_fin y frecuencia
    campos_permitidos = ['doctor', 'fecha_fin', 'frecuencia','descripcion']
    data_actualizada = {campo: valor for campo, valor in request.data.items() if campo in campos_permitidos}

    serializer = TratamientoActualSerializer(tratamiento, data=data_actualizada, partial=True)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

from datetime import date
@api_view(['PATCH'])
def finalizar_tratamiento(request, tratamiento_id):
    try:
        tratamiento = TratamientoActual.objects.get(id=tratamiento_id)
    except TratamientoActual.DoesNotExist:
        return Response({'error': 'No encontrado'}, status=404)
    
    tratamiento.finalizado = True
    tratamiento.fecha_fin =  date.today() 
    tratamiento.save()
    return Response({'mensaje': 'Tratamiento finalizado'})


from rest_framework.decorators import api_view, parser_classes


@api_view(['POST'])
@parser_classes([MultiPartParser, FormParser])
def agregar_seguimiento(request):
    serializer = SeguimientoTratamientoSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    print("Errores:", serializer.errors)  # Agregado para depuración
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# views.py


@api_view(['GET'])
def obtener_seguimientos(request, tratamiento_id):
    seguimientos = SeguimientoTratamiento.objects.filter(tratamiento_id=tratamiento_id).order_by('-fecha')
    serializer = SeguimientoTratamientoSerializer(seguimientos, many=True)
    return Response(serializer.data)




from .models import Medicamento
from .serializers import MedicamentoSerializer

@api_view(['GET'])
def listar_medicamentos(request):
    tipo = request.GET.get('tipo')
    if tipo:
        medicamentos = Medicamento.objects.filter(tipo__iexact=tipo)
    else:
        medicamentos = Medicamento.objects.all()
    
    serializer = MedicamentoSerializer(medicamentos, many=True)
    return Response(serializer.data)

from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from .models import Paciente, Usuario
from .serializers import PacienteSerializercedula

@api_view(['GET'])
def buscar_paciente_por_cedula(request):
    cedula = request.GET.get('cedula')

    if not cedula:
        return Response({"error": "Debe proporcionar una cédula."}, status=status.HTTP_400_BAD_REQUEST)

    try:
        usuario = Usuario.objects.get(cedula=cedula)
        paciente = Paciente.objects.get(id_usuario=usuario)
        serializer = PacienteSerializercedula(paciente)
        return Response(serializer.data, status=status.HTTP_200_OK)
    except Usuario.DoesNotExist:
        return Response({"error": "Usuario con esa cédula no encontrado."}, status=status.HTTP_404_NOT_FOUND)
    except Paciente.DoesNotExist:
        return Response({"error": "No hay paciente asociado a ese usuario."}, status=status.HTTP_404_NOT_FOUND)


from rest_framework.decorators import action
from django.utils.timezone import now
from .models import DoctorPaciente
from .serializers import DoctorPacienteSerializer

class DoctorPacienteViewSet(viewsets.ModelViewSet):
    queryset = DoctorPaciente.objects.all()
    serializer_class = DoctorPacienteSerializer

    def create(self, request, *args, **kwargs):
        data = request.data.copy()
        data['estado'] = 'pendiente'
        serializer = self.get_serializer(data=data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['post'])
    def aceptar(self, request, pk=None):
        relacion = self.get_object()
        if relacion.estado != 'pendiente':
            return Response({'detail': 'Ya fue respondida.'}, status=400)

        relacion.estado = 'aceptado'
        relacion.aprobado_en = now()
        relacion.save()
        return Response({'detail': 'Solicitud aceptada.'})

    @action(detail=True, methods=['post'])
    def rechazar(self, request, pk=None):
        relacion = self.get_object()
        if relacion.estado != 'pendiente':
            return Response({'detail': 'Ya fue respondida.'}, status=400)

        relacion.estado = 'rechazado'
        relacion.save()
        return Response({'detail': 'Solicitud rechazada.'})


class SolicitudesPorDoctorAPIView(APIView):
    def get(self, request, doctor_id):
        solicitudes = DoctorPaciente.objects.filter(doctor_id=doctor_id)
        serializer = DoctorPacienteDetalleSerializer(solicitudes, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)


class SolicitudesPorPacienteAPIView(APIView):
    def get(self, request, paciente_id):
        solicitudes = DoctorPaciente.objects.filter(paciente_id=paciente_id)
        serializer = DoctorPacienteDetalleSerializer(solicitudes, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)


from .serializers import DoctorSerializer

class DoctorPorUsuarioView(APIView):
    def get(self, request, id_usuario):
        try:
            doctor = Doctor.objects.get(id_usuario__id_usuario=id_usuario)
            serializer = DoctorSerializer(doctor)
            return Response(serializer.data, status=status.HTTP_200_OK)
        except Doctor.DoesNotExist:
            return Response({'error': 'Doctor no encontrado'}, status=status.HTTP_404_NOT_FOUND)

from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from .supabase_config import supabase

class SubirArchivoSupabase(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        archivo = request.FILES.get('archivo')

        if archivo:
            # Lee el archivo y lo sube al bucket correcto
            contenido = archivo.read()
            ruta = f"{archivo.name}"

            # Usa el bucket correcto: "examenes"
            supabase.storage.from_("examenes").upload(ruta, contenido, {"content-type": archivo.content_type})
            url_publica = supabase.storage.from_("examenes").get_public_url(ruta)

            return Response({
                "mensaje": "Archivo subido correctamente",
                "url": url_publica
            })
        
        return Response({"error": "No se envió archivo"}, status=400)
    


from rest_framework.views import APIView
from rest_framework.response import Response
import requests
import json

class HistoriaClinicaIAOpenRouter(APIView):
    def post(self, request):
        historia_clinica = request.data.get('historia_clinica')
        if not historia_clinica:
            return Response({'error': 'No se proporcionó la historia clínica'}, status=400)

        # Convertir a JSON string para incrustar en prompt
        historia_clinica_str = json.dumps(historia_clinica, indent=2, ensure_ascii=False)

        prompt = f"""
Eres un asistente médico experto en redacción clínica. A partir de la siguiente información estructurada en JSON sobre un paciente, genera un documento de historia clínica formal, claro y organizado, con apartados y subtítulos para facilitar la lectura médica:

{historia_clinica_str}

Organiza los datos en secciones como: Datos Personales, Signos Vitales, Consultas Médicas, Medicamentos Actuales, Alergias, Vacunas, Enfermedades Persistentes, Medicamentos Crónicos, Exámenes de Laboratorio y Exámenes de Imagenología.

Redacta en un lenguaje formal y profesional, usa listas o párrafos claros, e incluye fechas cuando estén disponibles.
"""

        try:
            response = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": "Bearer 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",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "meta-llama/llama-3-8b-instruct",
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.7
                }
            )
            data = response.json()
            resultado = data['choices'][0]['message']['content']
            return Response({
                "mensaje": "Historia clínica organizada generada con éxito",
                "historia_clinica_formateada": resultado
            })

        except Exception as e:
            return Response({'error': str(e)}, status=500)

from rest_framework import viewsets, permissions, serializers
from .models import SignosVitales, Paciente
from .serializers import SignosVitalesSerializer

class SignosVitalesViewSet(viewsets.ModelViewSet):
    serializer_class = SignosVitalesSerializer
    permission_classes = [permissions.AllowAny]

    def get_queryset(self):
        paciente_id = self.request.query_params.get('paciente_id')
        if paciente_id:
            return SignosVitales.objects.filter(paciente__id_paciente=paciente_id).order_by('-fecha')
        return SignosVitales.objects.all().order_by('-fecha')



    def perform_create(self, serializer):
        paciente_id = self.request.data.get('paciente')
        try:
            paciente = Paciente.objects.get(id_paciente=paciente_id)
        except Paciente.DoesNotExist:
            raise serializers.ValidationError("Paciente no válido.")
        
        serializer.save(paciente=paciente)

from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from usuarios.models import Paciente
import uuid

@api_view(['POST'])
def obtener_paciente_por_token(request):
    token = request.data.get('token')

    if not token:
        return Response({'error': 'Token no proporcionado'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        uuid_token = uuid.UUID(token)  # Verifica que sea un UUID válido
        paciente = Paciente.objects.get(token=uuid_token)
    except ValueError:
        return Response({'error': 'Token inválido'}, status=status.HTTP_400_BAD_REQUEST)
    except Paciente.DoesNotExist:
        return Response({'error': 'Paciente no encontrado'}, status=status.HTTP_404_NOT_FOUND)

    return Response({'paciente_id': paciente.id_paciente}, status=status.HTTP_200_OK)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import Consulta
from .serializers import ConsultaSerializer

class ConsultaListCreateView(APIView):
    def get(self, request):
        consultas = Consulta.objects.all()
        serializer = ConsultaSerializer(consultas, many=True)
        return Response(serializer.data)

    def post(self, request):
        serializer = ConsultaSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

from django.shortcuts import get_object_or_404

class ConsultaDetailView(APIView):
    def get(self, request, pk):
        consulta = get_object_or_404(Consulta, pk=pk)
        serializer = ConsultaSerializer(consulta)
        return Response(serializer.data)

    def put(self, request, pk):
        consulta = get_object_or_404(Consulta, pk=pk)
        serializer = ConsultaSerializer(consulta, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk):
        consulta = get_object_or_404(Consulta, pk=pk)
        consulta.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

from .models import DiagnosticoConsulta
from .serializers import DiagnosticoConsultaSerializer

class DiagnosticoConsultaView(APIView):
    def post(self, request):
        serializer = DiagnosticoConsultaSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class HistoriaClinicaPacienteView(APIView):
    def get(self, request, paciente_id):
        paciente = get_object_or_404(Paciente, id_paciente=paciente_id)
        serializer = HistoriaClinicaPacienteSerializer(paciente)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import ExamenFuncional
from .serializers import ExamenFuncionalSerializer
from django.shortcuts import get_object_or_404

class ExamenFuncionalListCreateView(APIView):
    def get(self, request):
        examenes = ExamenFuncional.objects.all()
        serializer = ExamenFuncionalSerializer(examenes, many=True)
        return Response(serializer.data)

    def post(self, request):
        serializer = ExamenFuncionalSerializer(data=request.data)
        if serializer.is_valid():
            # Opcional: validar que no exista examen para esta consulta si quieres limitar a uno por consulta
            consulta_id = serializer.validated_data['consulta'].id
            if ExamenFuncional.objects.filter(consulta_id=consulta_id).exists():
                return Response({"error": "Ya existe examen funcional para esta consulta."}, status=status.HTTP_400_BAD_REQUEST)
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class ExamenFuncionalDetailView(APIView):
    def get(self, request, pk):
        examen = get_object_or_404(ExamenFuncional, pk=pk)
        serializer = ExamenFuncionalSerializer(examen)
        return Response(serializer.data)

    def put(self, request, pk):
        examen = get_object_or_404(ExamenFuncional, pk=pk)
        serializer = ExamenFuncionalSerializer(examen, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk):
        examen = get_object_or_404(ExamenFuncional, pk=pk)
        examen.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


from .models import ExamenFisico
from .serializers import ExamenFisicoSerializer

class ExamenFisicoListCreateView(APIView):
    def get(self, request):
        examenes = ExamenFisico.objects.all()
        serializer = ExamenFisicoSerializer(examenes, many=True)
        return Response(serializer.data)

    def post(self, request):
        serializer = ExamenFisicoSerializer(data=request.data)
        if serializer.is_valid():
            consulta_id = serializer.validated_data['consulta'].id
            if ExamenFisico.objects.filter(consulta_id=consulta_id).exists():
                return Response({"error": "Ya existe examen físico para esta consulta."}, status=status.HTTP_400_BAD_REQUEST)
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class ExamenFisicoDetailView(APIView):
    def get(self, request, pk):
        examen = get_object_or_404(ExamenFisico, pk=pk)
        serializer = ExamenFisicoSerializer(examen)
        return Response(serializer.data)

    def put(self, request, pk):
        examen = get_object_or_404(ExamenFisico, pk=pk)
        serializer = ExamenFisicoSerializer(examen, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk):
        examen = get_object_or_404(ExamenFisico, pk=pk)
        examen.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.shortcuts import get_object_or_404
from django.utils.timezone import now, make_aware
from django.db.models import Count, Min, Max, Avg, Q, F, DurationField
from django.db.models.functions import TruncMonth, ExtractWeekDay
from django.core.exceptions import ValidationError
from datetime import datetime, timedelta
from django.utils import timezone
import logging

logger = logging.getLogger(__name__)

class EstadisticasDoctorView(APIView):
    def get(self, request, id_doctor):
        try:
            # Validaciones iniciales
            if not id_doctor:
                return Response(
                    {"error": "ID del doctor es requerido"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            try:
                id_doctor = int(id_doctor)
                if id_doctor <= 0:
                    raise ValueError
            except (ValueError, TypeError):
                return Response(
                    {"error": "ID del doctor debe ser un número entero positivo"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )

            doctor = get_object_or_404(Doctor, id_doctor=id_doctor)

            # Validar que el doctor esté activo (si tienes este campo)
            if hasattr(doctor, 'activo') and not doctor.activo:
                return Response(
                    {"error": "Doctor no está activo"}, 
                    status=status.HTTP_403_FORBIDDEN
                )

            hoy = now().date()
            # Usar timezone aware para las fechas
            inicio_mes = timezone.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            inicio_ano = timezone.now().replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
            hace_6_meses = timezone.now() - timedelta(days=180)

            # PACIENTES con validaciones mejoradas
            pacientes_qs = doctor.solicitudes_enviadas.filter(estado='aceptado')\
                .select_related('paciente__id_usuario')\
                .prefetch_related('paciente__perfil_bebe')

            total_pacientes = pacientes_qs.count()
            
            if total_pacientes == 0:
                logger.info(f"Doctor {id_doctor} no tiene pacientes")
            
            pacientes_nuevos_mes = pacientes_qs.filter(creado_en__gte=inicio_mes).count()
            pacientes_nuevos_ano = pacientes_qs.filter(creado_en__gte=inicio_ano).count()
            pacientes_activos_6_meses = pacientes_qs.filter(creado_en__gte=hace_6_meses).count()

            # Procesamiento de edades y demografía con validaciones
            edades = []
            distribucion_edad_sexo = {
                "niños_hombres": 0, "niñas_mujeres": 0,
                "jóvenes_hombres": 0, "jóvenes_mujeres": 0,
                "adultos_hombres": 0, "adultas_mujeres": 0,
                "adultos_mayores_hombres": 0, "adultas_mayores_mujeres": 0,
            }
            sexo_contador = {"masculino": 0, "femenino": 0, "otro": 0}
            pacientes_sin_fecha = 0
            pacientes_con_edad_invalida = 0

            for p in pacientes_qs:
                perfil = None
                sexo = None
                
                # Validar y obtener perfil
                if p.paciente and p.paciente.id_usuario and p.paciente.id_usuario.fecha_nacimiento:
                    perfil = p.paciente.id_usuario
                    sexo = (perfil.sexo or "").strip().lower()
                elif p.paciente and p.paciente.perfil_bebe and p.paciente.perfil_bebe.fecha_nacimiento:
                    perfil = p.paciente.perfil_bebe
                    sexo = (perfil.sexo or "").strip().lower()
                else:
                    pacientes_sin_fecha += 1
                    continue

                if perfil and perfil.fecha_nacimiento:
                    # Validar fecha de nacimiento
                    if perfil.fecha_nacimiento > hoy:
                        pacientes_con_edad_invalida += 1
                        logger.warning(f"Paciente con fecha de nacimiento futura: {perfil.fecha_nacimiento}")
                        continue
                    
                    edad = hoy.year - perfil.fecha_nacimiento.year - (
                        (hoy.month, hoy.day) < (perfil.fecha_nacimiento.month, perfil.fecha_nacimiento.day)
                    )
                    
                    # Validar edad razonable
                    if edad < 0 or edad > 150:
                        pacientes_con_edad_invalida += 1
                        logger.warning(f"Paciente con edad inválida: {edad}")
                        continue
                    
                    edades.append(edad)

                    # Normalizar sexo
                    if sexo in ['masculino', 'hombre', 'm', 'male']:
                        sexo = 'masculino'
                    elif sexo in ['femenino', 'mujer', 'f', 'female']:
                        sexo = 'femenino'
                    else:
                        sexo = 'otro'

                    # Clasificación edad+sexo
                    if edad < 13:
                        key = "niños_hombres" if sexo == "masculino" else "niñas_mujeres"
                    elif 13 <= edad <= 25:
                        key = "jóvenes_hombres" if sexo == "masculino" else "jóvenes_mujeres"
                    elif 26 <= edad <= 60:
                        key = "adultos_hombres" if sexo == "masculino" else "adultas_mujeres"
                    else:
                        key = "adultos_mayores_hombres" if sexo == "masculino" else "adultas_mayores_mujeres"
                    distribucion_edad_sexo[key] += 1

                    # Contador sexo general
                    sexo_contador[sexo] += 1

            # Estadísticas de edad mejoradas
            distribucion_edad = {
                "niños": sum(1 for e in edades if e < 13),
                "jóvenes": sum(1 for e in edades if 13 <= e <= 25),
                "adultos": sum(1 for e in edades if 26 <= e <= 60),
                "adultos_mayores": sum(1 for e in edades if e > 60),
            }
            
            edad_promedio = round(sum(edades) / len(edades), 1) if edades else 0
            edad_min = min(edades) if edades else 0
            edad_max = max(edades) if edades else 0

            distribucion_sexo = [{"sexo": k, "total": v} for k, v in sexo_contador.items()]

            # Pacientes por mes con validación
            pacientes_por_mes = (
                pacientes_qs.annotate(mes=TruncMonth('creado_en'))
                .values('mes')
                .annotate(total=Count('id'))
                .order_by('mes')
            )

            # TOP: Meses con más pacientes nuevos
            top_meses_pacientes = list(pacientes_por_mes.order_by('-total')[:3])

            # CONSULTAS con validaciones mejoradas
            consultas = doctor.consultas.all()
            total_consultas = consultas.count()
            
            if total_consultas == 0:
                logger.info(f"Doctor {id_doctor} no tiene consultas registradas")
            
            consultas_mes_actual = consultas.filter(fecha__gte=inicio_mes).count()
            consultas_ano_actual = consultas.filter(fecha__gte=inicio_ano).count()
            
            consultas_por_mes = (
                consultas.annotate(mes=TruncMonth('fecha'))
                .values('mes')
                .annotate(total=Count('id'))
                .order_by('mes')
            )
            
            consultas_por_dia = (
                consultas.annotate(dia=ExtractWeekDay('fecha'))
                .values('dia')
                .annotate(total=Count('id'))
            )

            # TOP: Días de la semana con más consultas
            top_dias_consultas = list(consultas_por_dia.order_by('-total')[:3])
            
            # Mapear números de día a nombres
            dias_nombres = {1: 'Domingo', 2: 'Lunes', 3: 'Martes', 4: 'Miércoles', 
                          5: 'Jueves', 6: 'Viernes', 7: 'Sábado'}
            for dia in top_dias_consultas:
                dia['nombre_dia'] = dias_nombres.get(dia['dia'], 'Desconocido')

            # Promedio de consultas con validación
            dias_registrados = consultas.aggregate(min_fecha=Min('fecha'), max_fecha=Max('fecha'))
            if dias_registrados['min_fecha'] and dias_registrados['max_fecha']:
                total_dias = (dias_registrados['max_fecha'].date() - dias_registrados['min_fecha'].date()).days + 1
                promedio_dia = round(total_consultas / total_dias, 2) if total_dias > 0 else 0
                promedio_semana = round(promedio_dia * 7, 2)
                promedio_mes = round(promedio_dia * 30, 2)
            else:
                promedio_dia = promedio_semana = promedio_mes = 0

            # DIAGNÓSTICOS con validaciones
            try:
                enfermedades_comunes = doctor.pacienteenfermedadcomun_set.values('enfermedad__tipo').annotate(total=Count('id')).order_by('-total')
                enfermedades_persistentes = doctor.pacienteenfermedadpersistente_set.values('enfermedad__tipo').annotate(total=Count('id')).order_by('-total')
                
                # TOP: Enfermedades más diagnosticadas
                top_enfermedades_comunes = list(enfermedades_comunes[:5])
                top_enfermedades_persistentes = list(enfermedades_persistentes[:5])
                
            except Exception as e:
                logger.error(f"Error obteniendo diagnósticos: {e}")
                enfermedades_comunes = []
                enfermedades_persistentes = []
                top_enfermedades_comunes = []
                top_enfermedades_persistentes = []

            # TRATAMIENTOS con validaciones mejoradas
            try:
                tratamientos = doctor.tratamientoactual_set.all()
                total_tratamientos = tratamientos.count()
                
                tratamientos_estado = tratamientos.values('finalizado').annotate(total=Count('id'))
                medicamentos_mas_recetados = tratamientos.values('medicamento__nombre_comercial').annotate(total=Count('id')).order_by('-total')[:10]
                
                # Calcular duración promedio de tratamientos finalizados de manera segura
                duracion_promedio_dias = None
                if hasattr(tratamientos.model, 'fecha_inicio') and hasattr(tratamientos.model, 'fecha_fin'):
                    # Usar diferencia en días para evitar problemas con AVG en fechas
                    tratamientos_con_duracion = tratamientos.filter(
                        finalizado=True,
                        fecha_inicio__isnull=False,
                        fecha_fin__isnull=False
                    ).annotate(
                        duracion_dias=F('fecha_fin') - F('fecha_inicio')
                    ).aggregate(
                        promedio_duracion=Avg('duracion_dias')
                    )
                    
                    if tratamientos_con_duracion['promedio_duracion']:
                        duracion_promedio_dias = tratamientos_con_duracion['promedio_duracion'].days
                
                # Efectividad de tratamientos (si tienes campo de resultado)
                tratamientos_exitosos = 0
                if hasattr(tratamientos.model, 'resultado_exitoso'):
                    tratamientos_exitosos = tratamientos.filter(resultado_exitoso=True).count()
                
                efectividad = round((tratamientos_exitosos / total_tratamientos) * 100, 2) if total_tratamientos > 0 else 0
                
            except Exception as e:
                logger.error(f"Error obteniendo tratamientos: {e}")
                tratamientos_estado = []
                medicamentos_mas_recetados = []
                duracion_promedio_dias = None
                efectividad = 0
                total_tratamientos = 0

            # Respuesta estructurada con validaciones
            response_data = {
                "doctor_id": id_doctor,
                "fecha_consulta": hoy.isoformat(),
                "validaciones": {
                    "pacientes_sin_fecha_nacimiento": pacientes_sin_fecha,
                    "pacientes_edad_invalida": pacientes_con_edad_invalida,
                    "total_pacientes_validos": len(edades),
                },
                "pacientes": {
                    "total": total_pacientes,
                    "nuevos_mes": pacientes_nuevos_mes,
                    "nuevos_ano": pacientes_nuevos_ano,
                    "activos_6_meses": pacientes_activos_6_meses,
                    "distribucion_edad": distribucion_edad,
                    "distribucion_edad_sexo": distribucion_edad_sexo,
                    "distribucion_sexo": distribucion_sexo,
                    "nuevos_por_mes": list(pacientes_por_mes),
                    "estadisticas_edad": {
                        "promedio": edad_promedio,
                        "minima": edad_min,
                        "maxima": edad_max,
                    },
                    "top_meses_nuevos": top_meses_pacientes,
                },
                "consultas": {
                    "total": total_consultas,
                    "mes_actual": consultas_mes_actual,
                    "ano_actual": consultas_ano_actual,
                    "por_mes": list(consultas_por_mes),
                    "por_dia_semana": list(consultas_por_dia),
                    "promedio_dia": promedio_dia,
                    "promedio_semana": promedio_semana,
                    "promedio_mes": promedio_mes,
                    "top_dias_semana": top_dias_consultas,
                },
                "diagnosticos": {
                    "comunes": list(enfermedades_comunes),
                    "persistentes": list(enfermedades_persistentes),
                    "top_comunes": top_enfermedades_comunes,
                    "top_persistentes": top_enfermedades_persistentes,
                },
                "tratamientos": {
                    "total": total_tratamientos,
                    "estado": list(tratamientos_estado),
                    "top_medicamentos": list(medicamentos_mas_recetados),
                    "efectividad_porcentaje": efectividad,
                    "duracion_promedio_dias": duracion_promedio_dias,
                }
            }

            return Response(response_data, status=status.HTTP_200_OK)

        except Doctor.DoesNotExist:
            return Response(
                {"error": "Doctor no encontrado"}, 
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            logger.error(f"Error inesperado en estadísticas del doctor {id_doctor}: {str(e)}")
            return Response(
                {"error": "Error interno del servidor"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )# Suponiendo que tienes el modelo Paciente y relaciones necesarias
from .models import Paciente
from .serializers import HistoriaClinicaPacienteSerializer  # tu serializer completo

from django.http import FileResponse, Http404
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions
from django.conf import settings
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile
from datetime import datetime
from collections import defaultdict

# Suponiendo que tienes el modelo Paciente y relaciones necesarias
from .models import Paciente
from .serializers import HistoriaClinicaPacienteSerializer

class DescargarHistoriaClinicaWord(APIView):
    permission_classes = [permissions.AllowAny]

    def get(self, request, id_paciente):
        try:
            paciente = Paciente.objects.get(id_paciente=id_paciente)
        except Paciente.DoesNotExist:
            return Response({"error": "Paciente no encontrado."}, status=404)

        # Serializamos todos los datos que usas en el JSON
        serializer = HistoriaClinicaPacienteSerializer(paciente)
        data = serializer.data

        # Crear archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            filepath = tmp.name
            self.generar_word(data, filepath)

        # Retornar el archivo para descarga
        try:
            response = FileResponse(open(filepath, 'rb'), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="historia_clinica_paciente_{id_paciente}.docx"'
            return response
        except Exception as e:
            return Response({"error": str(e)}, status=500)

    def format_date(self, date_string):
        """Formatea la fecha para mostrarla en español"""
        if not date_string:
            return "No registrado"
        try:
            if 'T' in str(date_string):
                date_obj = datetime.fromisoformat(str(date_string).replace('Z', '+00:00'))
            else:
                date_obj = datetime.strptime(str(date_string), '%Y-%m-%d')
            
            # Nombres de meses en español
            meses = [
                'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre'
            ]
            
            dia = date_obj.day
            mes = meses[date_obj.month - 1]
            año = date_obj.year
            
            return f"{dia} de {mes} de {año}"
        except:
            return "No registrado"

    def set_font_style(self, paragraph, font_name="Times New Roman", font_size=12):
        """Establece el estilo de fuente para un párrafo"""
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

    def add_field_with_label(self, paragraph, label, value, default_text="No registrado"):
        """Agrega un campo con etiqueta en negrita y valor"""
        if not value or (isinstance(value, str) and value.strip() == ''):
            value = default_text
        
        # Etiqueta en negrita
        run_label = paragraph.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.name = "Times New Roman"
        run_label.font.size = Pt(12)
        
        # Valor normal
        run_value = paragraph.add_run(str(value))
        run_value.font.name = "Times New Roman"
        run_value.font.size = Pt(12)
        return paragraph

    def add_page_break(self, doc):
        """Agrega un salto de página"""
        doc.add_page_break()

    def generar_seccion_consultas(self, doc, consultas_data):
        """Genera una sección completa y elegante de consultas médicas"""
        
        # Salto de página antes de consultas
        self.add_page_break(doc)
        
        heading = doc.add_heading('Consultas Médicas', level=1)
        self.set_font_style(heading, "Times New Roman", 14)
        
        if not consultas_data or len(consultas_data) == 0:
            p = doc.add_paragraph("No se registran consultas médicas para este paciente.")
            self.set_font_style(p)
            p.paragraph_format.space_after = Pt(12)
            return
        
        # Agregar resumen
        p_resumen = doc.add_paragraph()
        run_resumen = p_resumen.add_run(f"Total de consultas registradas: {len(consultas_data)}")
        run_resumen.bold = True
        run_resumen.font.name = "Times New Roman"
        run_resumen.font.size = Pt(12)
        p_resumen.paragraph_format.space_after = Pt(12)
        
        # Procesar cada consulta (más reciente primero, numeración descendente)
        total_consultas = len(consultas_data)
        for i, consulta in enumerate(consultas_data):
            numero_consulta = total_consultas - i  # Numeración descendente
            
            # Encabezado de consulta
            consulta_header = doc.add_heading(f'Consulta #{numero_consulta}', level=2)
            self.set_font_style(consulta_header, "Times New Roman", 13)
            consulta_header.paragraph_format.space_before = Pt(16)
            
            # Información básica
            p_info = doc.add_paragraph()
            self.add_field_with_label(p_info, "Fecha", self.format_date(consulta.get('fecha', '')))
            p_info.add_run('\n')
            self.add_field_with_label(p_info, "Motivo de Consulta", consulta.get('motivo', 'No registrado'))
            p_info.add_run('\n')
            self.add_field_with_label(p_info, "Síntomas", consulta.get('sintomas', 'No registrado'))
            p_info.paragraph_format.space_after = Pt(8)
            
            # Signos Vitales de la consulta
            signos_vitales = consulta.get('signos_vitales', [])
            p_signos = doc.add_paragraph()
            run_signos = p_signos.add_run("Signos Vitales:")
            run_signos.bold = True
            run_signos.underline = True
            run_signos.font.name = "Times New Roman"
            run_signos.font.size = Pt(12)
            p_signos.add_run('\n')
            
            if not signos_vitales:
                run_no_signos = p_signos.add_run("   No se registraron signos vitales en esta consulta")
                run_no_signos.font.name = "Times New Roman"
                run_no_signos.font.size = Pt(12)
            else:
                for signo in signos_vitales:
                    run_signo = p_signos.add_run(f"   {signo if signo else 'No registrado'}\n")
                    run_signo.font.name = "Times New Roman"
                    run_signo.font.size = Pt(12)
            p_signos.paragraph_format.space_after = Pt(8)
            
            # Examen Funcional
            examen_funcional = consulta.get('examen_funcional')
            if examen_funcional:
                p_func = doc.add_paragraph()
                run_func = p_func.add_run("Examen Funcional por Sistemas:")
                run_func.bold = True
                run_func.underline = True
                run_func.font.name = "Times New Roman"
                run_func.font.size = Pt(12)
                p_func.paragraph_format.space_after = Pt(4)
                
                sistemas = [
                    ('General', examen_funcional.get('general')),
                    ('Piel', examen_funcional.get('piel')),
                    ('Cabeza', examen_funcional.get('cabeza')),
                    ('Oídos', examen_funcional.get('oidos')),
                    ('Nariz', examen_funcional.get('nariz')),
                    ('Boca', examen_funcional.get('boca')),
                    ('Respiratorio', examen_funcional.get('respiratorio')),
                    ('Cardiovascular', examen_funcional.get('cardiovascular')),
                    ('Gastrointestinal', examen_funcional.get('gastrointestinal')),
                    ('Genitourinario', examen_funcional.get('genitourinario')),
                    ('Osteomuscular', examen_funcional.get('osteomuscular')),
                    ('Nervioso', examen_funcional.get('nervioso'))
                ]
                
                for sistema, evaluacion in sistemas:
                    if evaluacion and evaluacion.strip():
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.3)
                        self.add_field_with_label(p, f"{sistema}", evaluacion)
            else:
                p = doc.add_paragraph()
                run_no_func = p.add_run("Examen Funcional: No registrado")
                run_no_func.font.name = "Times New Roman"
                run_no_func.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(4)
            
            # Examen Físico
            examen_fisico = consulta.get('examen_fisico')
            if examen_fisico:
                p_fisico = doc.add_paragraph()
                run_fisico = p_fisico.add_run("Examen Físico:")
                run_fisico.bold = True
                run_fisico.underline = True
                run_fisico.font.name = "Times New Roman"
                run_fisico.font.size = Pt(12)
                p_fisico.paragraph_format.space_after = Pt(4)
                
                regiones = [
                    ('Aspecto General', examen_fisico.get('general')),
                    ('Piel', examen_fisico.get('piel')),
                    ('Uñas', examen_fisico.get('uñas')),
                    ('Cabeza', examen_fisico.get('cabeza')),
                    ('Ojos', examen_fisico.get('ojos')),
                    ('Nariz', examen_fisico.get('nariz')),
                    ('Oídos', examen_fisico.get('oidos')),
                    ('Boca y Faringe', examen_fisico.get('boca_faringe')),
                    ('Cuello', examen_fisico.get('cuello')),
                    ('Ganglios', examen_fisico.get('ganglios')),
                    ('Tórax', examen_fisico.get('torax')),
                    ('Pulmones', examen_fisico.get('pulmones')),
                    ('Corazón', examen_fisico.get('corazon')),
                    ('Abdomen', examen_fisico.get('abdomen')),
                    ('Genitales', examen_fisico.get('genitales')),
                    ('Recto', examen_fisico.get('recto')),
                    ('Osteomuscular', examen_fisico.get('osteomuscular')),
                    ('Neurológico/Psíquico', examen_fisico.get('neurologico_psiquico'))
                ]
                
                for region, hallazgo in regiones:
                    if hallazgo and hallazgo.strip():
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.3)
                        self.add_field_with_label(p, f"{region}", hallazgo)
            else:
                p = doc.add_paragraph()
                run_no_fisico = p.add_run("Examen Físico: No registrado")
                run_no_fisico.font.name = "Times New Roman"
                run_no_fisico.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(4)
            
            # Diagnóstico
            diagnostico = consulta.get('diagnostico')
            p_diag = doc.add_paragraph()
            run_diag = p_diag.add_run("Diagnóstico: ")
            run_diag.bold = True
            run_diag.font.name = "Times New Roman"
            run_diag.font.size = Pt(12)
            if diagnostico and diagnostico.get('descripcion'):
                run_valor = p_diag.add_run(diagnostico['descripcion'])
                run_valor.italic = True
                run_valor.font.name = "Times New Roman"
                run_valor.font.size = Pt(12)
            else:
                run_no_diag = p_diag.add_run("No registrado")
                run_no_diag.font.name = "Times New Roman"
                run_no_diag.font.size = Pt(12)
            p_diag.paragraph_format.space_after = Pt(8)
            
            # Tratamiento
            tratamientos = consulta.get('mensaje_tratamientos', '')
            p_trat = doc.add_paragraph()
            run_trat = p_trat.add_run("Tratamiento Prescrito:")
            run_trat.bold = True
            run_trat.underline = True
            run_trat.font.name = "Times New Roman"
            run_trat.font.size = Pt(12)
            p_trat.add_run('\n')
            
            if not tratamientos or tratamientos.strip() == '':
                run_no_trat = p_trat.add_run("   No registrado")
                run_no_trat.font.name = "Times New Roman"
                run_no_trat.font.size = Pt(12)
            else:
                # Separar múltiples tratamientos
                if '. ' in tratamientos:
                    tratamientos_lista = tratamientos.split('. ')
                else:
                    tratamientos_lista = [tratamientos]
                
                for j, tratamiento in enumerate(tratamientos_lista, 1):
                    if tratamiento.strip():
                        run_trat_item = p_trat.add_run(f"   {j}. {tratamiento.strip()}\n")
                        run_trat_item.font.name = "Times New Roman"
                        run_trat_item.font.size = Pt(12)
            p_trat.paragraph_format.space_after = Pt(8)
            
            # Observaciones
            observaciones = consulta.get('observaciones', '')
            p_obs = doc.add_paragraph()
            run_obs = p_obs.add_run("Observaciones y Recomendaciones: ")
            run_obs.bold = True
            run_obs.font.name = "Times New Roman"
            run_obs.font.size = Pt(12)
            if observaciones and observaciones.strip():
                run_obs_text = p_obs.add_run(observaciones)
                run_obs_text.font.name = "Times New Roman"
                run_obs_text.font.size = Pt(12)
            else:
                run_no_obs = p_obs.add_run("No registrado")
                run_no_obs.font.name = "Times New Roman"
                run_no_obs.font.size = Pt(12)
            
            # Línea separadora entre consultas (excepto la última)
            if i < len(consultas_data) - 1:
                separator = doc.add_paragraph("─" * 80)
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                separator.paragraph_format.space_before = Pt(12)
                separator.paragraph_format.space_after = Pt(12)
                self.set_font_style(separator)

    def generar_word(self, data, filepath):
        doc = Document()
        
        # Título principal
        title = doc.add_heading('Historia Clínica del Paciente', 0)
        self.set_font_style(title, "Times New Roman", 16)

        # Datos del Paciente
        heading1 = doc.add_heading('Datos del Paciente', level=1)
        self.set_font_style(heading1, "Times New Roman", 14)
        
        p1 = doc.add_paragraph(f"Nombre: {data.get('nombre', 'No registrado')} {data.get('apellido', '')}")
        self.set_font_style(p1)
        p2 = doc.add_paragraph(f"Tipo de Sangre: {data.get('tipo_sangre', 'No registrado')}")
        self.set_font_style(p2)
        p3 = doc.add_paragraph(f"Cédula de identidad: {data.get('cedula', 'No registrado')}")
        self.set_font_style(p3)
        p4 = doc.add_paragraph(f"Edad: {data.get('edad', 'No registrado')}")
        self.set_font_style(p4)
        p5 = doc.add_paragraph(f"Sexo: {data.get('sexo', 'No registrado')}")
        self.set_font_style(p5)
        p6 = doc.add_paragraph(f"Número de teléfono: {data.get('telefono', 'No registrado')}")
        self.set_font_style(p6)
        p7 = doc.add_paragraph(f"Nacionalidad: {data.get('nacionalidad', 'No registrado')}")
        self.set_font_style(p7)

        # Signos Vitales
        sv = data.get('signos_vitales')
        heading2 = doc.add_heading('Última toma de Signos Vitales', level=1)
        self.set_font_style(heading2, "Times New Roman", 14)
        if sv:
            p_sv1 = doc.add_paragraph(f"Fecha: {self.format_date(sv.get('fecha', ''))}")
            self.set_font_style(p_sv1)
            p_sv2 = doc.add_paragraph(f"Peso: {sv.get('peso', 'No registrado')} kg")
            self.set_font_style(p_sv2)
            p_sv3 = doc.add_paragraph(f"Altura: {sv.get('altura', 'No registrado')} m")
            self.set_font_style(p_sv3)
            p_sv4 = doc.add_paragraph(f"IMC: {sv.get('imc', 'No registrado')}")
            self.set_font_style(p_sv4)
            p_sv5 = doc.add_paragraph(f"Presión arterial: {sv.get('presion_sistolica', 'No registrado')}/{sv.get('presion_diastolica', 'No registrado')} mmHg")
            self.set_font_style(p_sv5)
            p_sv6 = doc.add_paragraph(f"Frecuencia cardíaca: {sv.get('frecuencia_cardiaca', 'No registrado')} lpm")
            self.set_font_style(p_sv6)
            p_sv7 = doc.add_paragraph(f"Temperatura: {sv.get('temperatura', 'No registrado')} °C")
            self.set_font_style(p_sv7)
            p_sv8 = doc.add_paragraph(f"Glucosa: {sv.get('glucosa', 'No registrado')} mg/dL")
            self.set_font_style(p_sv8)
            p_sv9 = doc.add_paragraph(f"Saturación O₂: {sv.get('spo2', 'No registrado')}%")
            self.set_font_style(p_sv9)
            p_sv10 = doc.add_paragraph(f"Observaciones: {sv.get('observaciones', 'No registrado')}")
            self.set_font_style(p_sv10)
        else:
            p_no_sv = doc.add_paragraph("No se han registrado signos vitales para este paciente.")
            self.set_font_style(p_no_sv)

        # NUEVA PÁGINA - Tratamientos Actuales
        self.add_page_break(doc)
        heading3 = doc.add_heading("Tratamientos Actuales", level=1)
        self.set_font_style(heading3, "Times New Roman", 14)
        tratamientos_actuales = data.get('tratamientos_actuales', [])
        if tratamientos_actuales:
            for t in tratamientos_actuales:
                texto = t.get('texto_formateado')
                if texto:
                    parrafo = doc.add_paragraph()
                    for i, line in enumerate(texto.split('\n')):
                        run = parrafo.add_run(line + '\n')
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        if i == 0:
                            run.bold = True
        else:
            p_no_trat = doc.add_paragraph("No se registran tratamientos actuales.")
            self.set_font_style(p_no_trat)

        # NUEVA PÁGINA - Alergias conocidas
        self.add_page_break(doc)
        heading4 = doc.add_heading('Alergias conocidas', level=1)
        self.set_font_style(heading4, "Times New Roman", 14)
        alergias = data.get('alergias', [])
        if alergias:
            for a in alergias:
                nombre = a.get('nombre_alergia', 'No registrado')
                tipo = a.get('tipo_alergia', 'No registrado')
                gravedad = a.get('gravedad', 'No registrado')
                observacion = a.get('observacion', '')

                p = doc.add_paragraph()
                run = p.add_run(f"{nombre} ({tipo})\n")
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

                run2 = p.add_run(f"Gravedad: {gravedad}\n")
                run2.font.name = "Times New Roman"
                run2.font.size = Pt(12)

                if observacion and observacion.strip():
                    run3 = p.add_run(f"Observación: {observacion}\n")
                    run3.font.name = "Times New Roman"
                    run3.font.size = Pt(12)
                else:
                    run3 = p.add_run(f"Observación: No registrado\n")
                    run3.font.name = "Times New Roman"
                    run3.font.size = Pt(12)

                p.paragraph_format.space_after = Pt(8)
        else:
            p_no_alerg = doc.add_paragraph("No se registran alergias conocidas.")
            self.set_font_style(p_no_alerg)

        # NUEVA PÁGINA - Enfermedades Preexistentes
        self.add_page_break(doc)
        heading5 = doc.add_heading('Enfermedades Preexistentes', level=1)
        self.set_font_style(heading5, "Times New Roman", 14)
        enfermedades = data.get('enfermedades_persistentes', [])
        if enfermedades:
            for e in enfermedades:
                nombre = e.get('nombre', 'No registrado')
                tipo = e.get('tipo', 'No registrado')
                fecha = e.get('fecha_diagnostico', '')
                observacion = e.get('observacion', '')

                p = doc.add_paragraph()
                run = p.add_run(f"{nombre} ({tipo})\n")
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

                run2 = p.add_run(f"Diagnóstico desde: {self.format_date(fecha)}\n")
                run2.font.name = "Times New Roman"
                run2.font.size = Pt(12)

                if observacion and observacion.strip():
                    run3 = p.add_run(f"Observación: {observacion}\n")
                    run3.font.name = "Times New Roman"
                    run3.font.size = Pt(12)
                else:
                    run3 = p.add_run(f"Observación: No registrado\n")
                    run3.font.name = "Times New Roman"
                    run3.font.size = Pt(12)

                p.paragraph_format.space_after = Pt(8)
        else:
            p_no_enf = doc.add_paragraph("No se registran enfermedades preexistentes.")
            self.set_font_style(p_no_enf)

        # NUEVA PÁGINA - Vacunas
        self.add_page_break(doc)
        heading6 = doc.add_heading('Vacunas', level=1)
        self.set_font_style(heading6, "Times New Roman", 14)
        vacunas = data.get('vacunas', [])
        if vacunas:
            vacunas_agrupadas = defaultdict(list)
            for v in vacunas:
                nombre = v.get('nombre_vacuna', 'No registrado')
                dosis = v.get('dosis', 'No registrado')
                fecha = v.get('fecha_aplicacion', '')
                observacion = v.get('observacion', '')
                vacunas_agrupadas[nombre].append({
                    'dosis': dosis,
                    'fecha': fecha,
                    'observacion': observacion
                })

            for nombre, detalles in vacunas_agrupadas.items():
                p = doc.add_paragraph()
                run = p.add_run(f"{nombre}\n")
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

                for det in detalles:
                    dosis = det['dosis']
                    fecha = self.format_date(det['fecha'])
                    run2 = p.add_run(f"Dosis: {dosis} – Fecha: {fecha}\n")
                    run2.font.name = "Times New Roman"
                    run2.font.size = Pt(12)

                observaciones = {d['observacion'] for d in detalles if d['observacion'] and d['observacion'].strip()}
                if observaciones:
                    run3 = p.add_run(f"Observación: {', '.join(observaciones)}\n")
                    run3.font.name = "Times New Roman"
                    run3.font.size = Pt(12)
                else:
                    run3 = p.add_run(f"Observación: No registrado\n")
                    run3.font.name = "Times New Roman"
                    run3.font.size = Pt(12)

                p.paragraph_format.space_after = Pt(8)
        else:
            p_no_vac = doc.add_paragraph("No se registran vacunas aplicadas.")
            self.set_font_style(p_no_vac)

        # NUEVA PÁGINA - Medicamentos Frecuentes
        self.add_page_break(doc)
        heading7 = doc.add_heading('Medicamentos Frecuentes', level=1)
        self.set_font_style(heading7, "Times New Roman", 14)
        medicamentos = data.get('medicamentos_cronicos', [])
        if medicamentos:
            for m in medicamentos:
                p = doc.add_paragraph()
                run = p.add_run(f"{m.get('nombre', 'No registrado')}\n")
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                
                dosis = m.get('dosis', '')
                if dosis and dosis.strip():
                    run2 = p.add_run(f"Dosis: {dosis}\n")
                else:
                    run2 = p.add_run(f"Dosis: No registrado\n")
                run2.font.name = "Times New Roman"
                run2.font.size = Pt(12)
                
                frecuencia = m.get('frecuencia', '')
                if frecuencia and frecuencia.strip():
                    run3 = p.add_run(f"Frecuencia: {frecuencia}\n")
                else:
                    run3 = p.add_run(f"Frecuencia: No registrado\n")
                run3.font.name = "Times New Roman"
                run3.font.size = Pt(12)
                
                observaciones = m.get('observaciones', '')
                if observaciones and observaciones.strip():
                    run4 = p.add_run(f"Observación: {observaciones}\n")
                else:
                    run4 = p.add_run(f"Observación: No registrado\n")
                run4.font.name = "Times New Roman"
                run4.font.size = Pt(12)
                
                p.paragraph_format.space_after = Pt(8)
        else:
            p_no_med = doc.add_paragraph("No se registran medicamentos frecuentes.")
            self.set_font_style(p_no_med)

        # NUEVA PÁGINA - Exámenes de Laboratorio
        self.add_page_break(doc)

        # Exámenes de Laboratorio
        heading8 = doc.add_heading('Exámenes de Laboratorio', level=1)
        self.set_font_style(heading8, "Times New Roman", 14)
        examenes_lab = data.get('examenes_laboratorio', [])
        if examenes_lab:
            for e in examenes_lab:
                p = doc.add_paragraph()
                run = p.add_run(f"{e.get('nombre_examen', 'Examen no especificado')}\n")
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run2 = p.add_run(f"Categoría: {e.get('categoria', 'No especificada')} – Fecha: {self.format_date(e.get('fecha_realizacion', ''))}\n")
                run2.font.name = "Times New Roman"
                run2.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(8)
        else:
            p_no_lab = doc.add_paragraph("No se registran exámenes de laboratorio.")
            self.set_font_style(p_no_lab)

        # Imagenología
        self.add_page_break(doc)
        heading9 = doc.add_heading('Imagenología', level=1)
        self.set_font_style(heading9, "Times New Roman", 14)
        examenes_img = data.get('examenes_imagenologia', [])
        if examenes_img:
            for e in examenes_img:
                p = doc.add_paragraph()
                run = p.add_run(f"{e.get('nombre_examen', 'Examen no especificado')}\n")
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run2 = p.add_run(f"Categoría: {e.get('categoria', 'No especificada')} – Fecha: {self.format_date(e.get('fecha_realizacion', ''))}\n")
                run2.font.name = "Times New Roman"
                run2.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(8)
        else:
            p_no_img = doc.add_paragraph("No se registran estudios de imagenología.")
            self.set_font_style(p_no_img)

        # Consultas Médicas - AHORA AL FINAL
        self.generar_seccion_consultas(doc, data.get('consultas', []))

        doc.save(filepath)



